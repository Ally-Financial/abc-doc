'''
 * Copyright 2021 Ally Financial, Inc.
 *
 * Licensed under the Apache License, Version 2.0 (the "License"); you may not use this file except
 * in compliance with the License. You may obtain a copy of the License at
 *
 * http://www.apache.org/licenses/LICENSE-2.0
 *
 * Unless required by applicable law or agreed to in writing, software distributed under the License
 * is distributed on an "AS IS" BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express
 * or implied. See the License for the specific language governing permissions and limitations under
 * the License.
 '''
#library to capture Mouse Events
from pynput.mouse import Listener
from pynput import mouse
import logging
#library to capture Screenshot
import pyautogui
#library to prepare and work on word File
from docx import Document
from docx.shared import Inches
#library to Build Application GUI
import PySimpleGUI as sg
#library to get and Set system details
import os
from datetime import datetime
#library to work on images and merge 2 images
from PIL import Image, ImageDraw
#import pii_DetectBlur_Image
#nexus 
v_enable_logging = True
# for padding white spaces for better layout
vSzInput=(3,10)
vSzButton=(25,3)
vSzIcon=(3,3)
# parameters to be changed for individual or organizaiton needs
vFinalFilenameStartWith = 'Ally - '
vFilename = 'Ally - '
vContinueProcessCapture = False
# parameters to be changed for individual or organizaiton needs
vScreenTextFont = 'Calibri 12'
vDesktopPath = (os.environ['USERPROFILE'] + '\Desktop')
# convert images or logo to 64 bit string for easy portability
vLogoOnScreen = b'iVBORw0KGgoAAAANSUhEUgAAADIAAAAtCAYAAADsvzj/AAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAAFuoAABbqAeWOQxAAAAbCSURBVGhD7VlpbFRVFB6KSFQWi2Aps7xlBjAoEG2MGIlEZakUqW1nee9u0xaD+gPj8kfi0kQQBNJEgtEQBU0MEreEaExUjBL+mBgDBmIBg8aYKEESRHahMH7n9RaG8dF5M51WTfiS023ePfd8557tvoau4Ar+J0iFUkNT41Ij3Fp3bEOkISxrZUzaMsbDPJI1suMbYg3VT0SS1+jH/1sgw+UEOckJi3uZLVq5zZ+XllwrrexGfH8HRN6VdvZtZan1+H0VN/iTwsg+yCPp23mM13aEOqq0qn8H8+HhTDQzV5ryOWGKrdySv3JLnMnGW3PtiUU9Er8oixIP5VrjbTlhyZy01J+Q3dyUb4JYe9pMT09GZgzuSbHRrNqJ8QYY8QaM+o2Ma4u354iAsrM5YcuiQs/RGiJLxHCSXzOLPdUSdW7GFgN/QkmLTeW2XAPPH+g13s/QUoV0ybjCz3wbt5Wbqk6N1ltWHEOZwe7npviSvNlaIQKF0pYAIVsd5hZflrJTMb13ZVAXqhsmDOEIW+zzYtzHgEoKnTLI5JB3G1zbnajN6D+4oVwk8k9Ze2BOwU9UPJtTVjaHPNyYjAlLm1I+kpYzByR2ZxODR6JXPDIIY5TsNXKkvEGbVDpQWuPCUp9WKqHLEb33EW7yh+tCi4dp00rCEGmKF5WtTpJnCjcYTGlNeL1nJ4uxu7RtweFYYg63RVc5p0GJSuS9hLXlecg5EgqTvL+XJEQGTu1sNBqv1yYWx5TQlKt5XLwGBTAg+KYeARiLnDqO3/cgUT93Tb6JWXwjyukmlO6t+LwLz52gZ0shpCeC/SIq5mgziyNluzMxPuwq5TSomWHNWSTmd8wSKxEGs1VEhRfruK4P1Q9vCjdFMlE+FyPJGhi1C9Ld0wT9dRZKjz1qOR9TP8oztBhQv5+Ft44G9Zj27kkMi++lDT5Dq+kT6Qi7mxtiC8icggN89RaK1/1ttZ1F2FSt5vJIjm8bh0n1fQoRP2WFokl0M1NtaJ/sTNBqAiFTkzGlyTcLU54hPX7688Xby1LHHEsltYrLI2nIO1FydwQNK2qSzJKfuUbzTVpFSUiH09PhtG1BiJCQXcxmHXW1dddqFf5A/LpZW/0chAiRQGX7I2MwppeXBfSIJQixQ0HyhaoXisnmTE2jqZf7w7X50/DOkSAeopjlaJhJMzlZLy8LNLojvLYHCWeqXtJUXzlhZ5pe7g9mik54J1DMtmJSxfMrRY24Ti8vF1VUooMQIefhuR0Iyb6Limu4uI4Gi1fyjmPwx/XSfoEbkkryuWIOpJDHM3uo6uml/mCGeJ2mTj8lhUIe5DG5RC/tF1hMLIcD/wpCBPvuUwa7Ry/1BxJvHTzjlTo/RflCx+wa/Jm6uvKGuXwgRAMTESCCS17fRNCVX8A19mQQIj3DnHprYbS9pP7hh4oTQVdfDGUHgpRCrXQvM1TfSgOg4kQQWnORI11Br7Q9lUZ0No4uYSr1QcWJNIebJ0LhF0FKIYk3lpvqIKpOO71l1GpKRsWJhEIdVa7F10JpoF4COU+nh267H/m1aHb17LJe45RStQISwaXKSLnKkr+UcjP0Om4cJ2PJVY7p3EHvfmeFZl2lVRaFG+MrKk4kGUmGpSk/ChpevUKbeG9aLPU9t+U6umezKGvkCdxBfMSNu/Nc252vopk4NURUy9MVJUJAh38UHjpYzn2dTkffs7ulLQ/j+++kq1Dw96OYnPeKmGhBOFNnL1r2SybCxzbVwkMflkMkX2i9RyxP6NTopTam7BMCfYvykpmsA88XzcuSiRBwXV2gLLWn1BArJkQO9/ljSPCO1LhZI7y9Kl+1LkGVMMRSxPuR/p5Mr5AeGHuCGfyllH3xRfVAEwklRyXH4LL1CuL3dCXCDMaehtGdbu2CsXoLDwNOhIAKE5EWX4/ELftlnbcO9xyE1KuO5dRo1RcwGESG0JcmJD8S82Vly0OUM8U2zBea2zD2dCtbbeBhHvG0FmBQTgTwyCycvHAkhsrHsOFOSLeO939sli/683PCVJuaY8z2tPmgpM5uqrKJXAIZdWfCgxsQKvux+Rkqpz3T8qVGkFGemOoDOSEzSS/3hWvyFTjps2TohXUFQvq9d1uW+oHH+H16af9Qn6gf7kbdBzhyByHzDUKHGtzZ3k17+wW897GIiFv0sssCBWUZdBwnx8AJpy4nPaGluipyIvmgibfFapkmDfkI5qxXuSk+ESb/Ft79EYZtScf4bfrRPiEjch5um6scm3WyOFvtL87qDAbatOUszSRUXC8dGIiaphudhHOrY4j6ZE2y//9pCoRQ6G8oxV0jL1garQAAAABJRU5ErkJggg=='
vAppLogo = b'iVBORw0KGgoAAAANSUhEUgAAAEgAAABICAMAAABiM0N1AAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAzUExURQAAACCPjxiXhxuVihyXhx2WiRuXihuWiRyXiR2YixuYihyYihyYihyYihuYihyYihyYinvhKU4AAAAQdFJOUwAQIDBAUGBwgI+fr7/P3+8jGoKKAAAACXBIWXMAAA7DAAAOwwHHb6hkAAADHUlEQVRYR+2X25asIAxERfCCN/z/rz0VSBAEWvusNW+9H6btFiNJVYDpfvz4e9Q88+fS+4v/ZTvPHRHUep4H//Qlyv+dTuCm8aDP6fr9Pdqd69gZRwEEZ5RZ3Gl4yDuQSouVh7zD51LH8ZBXaH6oysvclDG9+ZAZ5DNaDzy6jcrqC459tXbdilyfbDXwOI/bJs2/d2qCpxKe5pRUx83BNErL22e+QzwF6nnceW6YjDLzHnI9LNW4X/wX4qnkMiM3UjZ5Zcjahq+fAoleDjM3e7i+OBAdnvfssXolGg0QcH2nLF9nWMwzTnNtzSoqj/lQv9fYVKpHvYUV3z3PsRkHkZI6tcwkkqAtq3kFcDcZWIXN6HRuyzvIO141CLdtpwq9PFxDtD87s9UmoYpQzC+MBUfPCc8ypYZqY3gjEq+uRgfeHyJhSlwlv/reiDrNiSoJFEfWhqEb/SdsWRogdlGf9aYQ4nDtlssqS3g6QQTHrlPJLItDucll2PYy+NYu4mF87LY8DjmaF6dKapL/ElcS+ImfvcdB+lz2aueGAtrYS2Q3/3QRB7qHQJXEQHCPTWfkny/jIFAQpCZ/14car5ciIdJaxkFqQWSYt0B8hO4OFyCW4BYHxebRkOaOmAfSXvJzpHucxCJtH0VFPFKnHBhSFsFyJYmjp2wRQaQiDvSUIXR4KuCevzXt5cqIkxK5qb7WhmbFsPoyEsEyIs3bwN8ll1W67QJvYmXq87mW2v7zUjt3PZca21wVWZEgaVxUSrC1iazwXBWRlHLPjx4J6HdeSkFNMxiA7wJo3ogEvWXT9kWoYuLDqGdqywurLlu5raka6s2jyNG3swhw6HbNv+66pVkgbuz0ttxP3n7x6P1hNp7rhGCpAIPlB91qEEbJVyrjZ5KjBglN9DjF+jSSMM8ntnxPw78RghrWJMxzaokHAvu6zNZufJSMtJRPGIwW7ers2gzmOY7nOttWeEoq45ZICpn1PR+atnVMqwPxtknlErpBDZs/OH+FN5zffN08+Uz9tqpeVvkO5KP1vafPr2pzR1n0OzH7DffHjz+l6/4Bmz911Xhw3/QAAAAASUVORK5CYII='

def add_to_log(text):
    if v_enable_logging:
        print('Timestamp '+str(docnow.strftime("%m-%d-%Y %H%M")))
        print(str(docnow.strftime("%m-%d-%Y %H%M%S")) + ' : ' + text)
        logging.info(text)

# Add your new theme colors and settings
sg.LOOK_AND_FEEL_TABLE['Ally'] = {'BACKGROUND': '#FFFFFF',
                            'TEXT': '#000000',
                            'INPUT': '#FFF7F0',
                            'TEXT_INPUT': '#000000',
                            'SCROLL': '#FFF7F0', 
                            'BUTTON': ('#FFFFFF' , '#8A3575'),
                            'PROGRESS': "#FFF7F0", 
                            'SCROLL': '#FFF7F0', 
                            'BORDER': 1,
                            'SLIDER_DEPTH':0, 
                            'PROGRESS_DEPTH':0
                            }

sg.ChangeLookAndFeel('Ally')

column_input_padding = [[sg.Text('',  size=vSzInput)]]
column_button_padding = [[sg.Text('',  size=vSzButton)]]
column_icon_padding = [[sg.Text('',  size=vSzIcon)]]

column_input = [[sg.Text('Save Document here', size=(20, 1), justification='left'),
     sg.InputText(vDesktopPath), sg.FolderBrowse( size=(8,1))],
    [sg.Text('Process Name ', size=(20, 1), justification='left'),sg.InputText('')],
    [sg.Text('Process Description ', size=(20, 1), justification='left'),
     sg.MLine(default_text='', size=(45, 3),no_scrollbar = True)],
    [sg.Text('Systems Impacted ', size=(20, 1), justification='left'),
     sg.InputText(default_text='', size=(45, 3))]]

column_button = [[sg.Button("Start", size =(8,1)),
                sg.Button("Pause",disabled=True, size =(8,1)),
                sg.Button("Continue",disabled=True, size =(8,1)),
                sg.Button("Stop",disabled=True, size =(8,1))]]

column_ally_logo = [[sg.Image(data=vLogoOnScreen)]]
#Preparing layout with input fields and button 
layout = [[
            [
            sg.Column(column_input_padding, element_justification='c' ), 
            sg.Column(column_input, element_justification='l')],
            [sg.Column(column_button_padding, element_justification='c' ), 
            sg.Column(column_button, element_justification='l')],
            [sg.Column(column_ally_logo, element_justification='l')
            ]
          ]]

#Function to capture mouse Move
def on_move(x, y):
    pass 

#Function to paste image 2(fg_img) on image 1(bg_img)
# alpha represent opacity 
# Box represent position on Image 1 where image 2 will be added
def trans_paste(fg_img,bg_img,alpha=1.0,box=(0,0)):
    fg_img_trans = Image.new("RGBA",fg_img.size)
    fg_img_trans = Image.blend(fg_img_trans,fg_img,alpha)
    bg_img.paste(fg_img_trans,box,fg_img_trans)
    return bg_img

#Function to prepare word file with inputs provided by users
def prepare_word_file(document,values):
    try:
            add_to_log('Started preparing word file in Current temaplet folder')
            document.add_heading('Process Name', level=1)
            document.add_paragraph(values[1])
            #document.add_page_break()

            document.add_heading('Process Description', level=1)
            document.add_paragraph(values[2])
            #document.add_page_break()
            
            document.add_heading('System Impacted', level=1)
            document.add_paragraph(values[3])
            #document.add_page_break()
            
            document.add_heading('As-Is Process', level=1)
            document.add_paragraph('Process screenshot will be added below')
    except Exception as e:
        add_to_log(' Error while preparing Word file '+ str(e))

#Function to capture Screenshot and Add to word file
def capture_and_save(x, y):
        add_to_log('Mouse Clicked ')
        myScreenshot = pyautogui.screenshot()
        now = datetime.now()
        timestamp = datetime.timestamp(now)
        path_to_save=vFilePath+r'/file_name'+str(timestamp)+'.png'
        path_to_save_bg=vFilePath+r'/file_name_bg'+str(timestamp)+'.png'
        path_to_save_fg=vFilePath+r'/file_name_fg'+str(timestamp)+'.png'
        myScreenshot.save(path_to_save)
        add_to_log('Screenshot Saved ')
        #path_to_save_ret = pii_DetectBlur_Image.main(path_to_save)
        path_to_save_ret = path_to_save

        bg_img = Image.open(path_to_save_ret)
        fg_img = Image.open('Mouse Icon.png')
        add_to_log('Screenshot Saved ')
        p = trans_paste(fg_img.convert('RGBA'),bg_img.convert('RGBA'),1,(x,y))
        #p = p.crop((left, top, right, bottom))
        p.save(path_to_save_bg)
        add_to_log('Image Saved ')

        document.add_picture(path_to_save_bg, width=Inches(6.25))
        document.save(vFilePath+r'/'+vFilename)
        add_to_log('Image Added to file ')    
        if os.path.exists(path_to_save):
            os.remove(path_to_save)
        if os.path.exists(path_to_save_ret):
            os.remove(path_to_save_ret)
        if os.path.exists(path_to_save_bg):
            os.remove(path_to_save_bg)
        add_to_log('Image deleted ')

#Function to capture on click event and add it to screenshot 
def on_click(x, y, button, pressed):
    if pressed and vContinueProcessCapture:
        capture_and_save(x, y)

#Function to skip scroll event( can be extended in Future)       
def on_scroll(x, y, dx, dy):
    pass 

#Create Window and add fetch input
window = sg.Window(' ABC Doc - Automated Business process Capture and Documentation ', 
    layout,
    default_element_size=(45, 22), 
    grab_anywhere=False,
    icon = vAppLogo,
    progress_bar_color='red',
    font = vScreenTextFont,
    element_padding=5
    )

    #Create Base Documentcls

try:
    document = Document()
    document.add_heading('Process Definition Document (PDD)', 0)
    docnow = datetime.now()
    doctimestamp = docnow.strftime("%m-%d-%Y %H%M")
    vFilename = vFilename+' '+ str(doctimestamp)+'.docx'
    vFilePrepared = False
except  Exception as e:
    add_to_log(' Error while getting date , time and preparing file name '+ str(e))

# Create window exists, fetch events on window
while True:
    try:
        #Create Base Document
        event, values = window.read()
        vFilePath = values[0]
        add_to_log(str(values )) 
        if vFilePrepared == False:
            add_to_log(' 3 ')
            prepare_word_file(document,values)
            add_to_log(' 4 ')
            vFilePrepared = True
            add_to_log(' 5 ')
    except Exception as e:
        add_to_log(' Error while Preparing window '+ str(e))

    # Steps to be performed on Click "Start"
    if event == "Start":
        vContinueProcessCapture=True
        window.Minimize()
        add_to_log('Start Clicked')
        listener = mouse.Listener(on_move=on_move,on_click=on_click,on_scroll=on_scroll)
        listener.start()
        window.find_element("Pause").Update(disabled=False)
        window.find_element("Stop").Update(disabled=False)
        window.find_element("Start").Update(disabled=True)
        window.find_element("Continue").Update(disabled=True)
        window.Refresh()

    # Steps to be performed on Click "Stop or Window Closed"    
    if event == "Stop" or event == sg.WIN_CLOSED:
        try:
            os.rename(vFilePath+r'/'+vFilename,vFinalFilenameStartWith+values[1]+r'/'+' '+str(doctimestamp)+'.docx') 
            add_to_log('Stop Clicked')#listener.stop()
            break
        except:
            break     
    
    if event == "Pause":
        vContinueProcessCapture=False
        add_to_log('Pause Clicked')
        window.find_element("Continue").Update(disabled=False)
        window.find_element("Pause").Update(disabled=True)
        window.Refresh()
    
    if event == "Continue":
        vContinueProcessCapture=True
        add_to_log('Continue Clicked')
        window.find_element("Continue").Update(disabled=True)
        window.find_element("Pause").Update(disabled=False)
        window.Refresh()

window.close()
